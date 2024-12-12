import tkinter as tk
from tkinter import ttk, PhotoImage, messagebox
from docx import Document
from interface.login.pagina_aluno.menus_aluno.DadosUsuario import abrir_dados_usuario
from interface.login.pagina_aluno.menus_aluno.Descobrir import descobrir
from interface.login.pagina_aluno.menus_aluno.MeusCursos import meusCursos
from util.db import conexaoBanco
from datetime import datetime
import os

def abrir_curso(tree):
    """
    Abre os detalhes do curso selecionado no Treeview.
    """
    try:
        # Obter o item selecionado no Treeview
        selected_item = tree.focus()
        if not selected_item:
            messagebox.showwarning("Seleção Inválida", "Por favor, selecione um curso.")
            return

        # Obter os valores da linha selecionada
        valores = tree.item(selected_item, "values")
        if not valores:
            messagebox.showerror("Erro", "Erro ao obter os dados do curso.")
            return

        # Extrair as informações do curso
        titulo_curso, nota, presenca, certificado = valores

        # Criar uma nova janela para exibir os detalhes do curso
        detalhes_janela = tk.Toplevel()
        detalhes_janela.title(f"Detalhes do Curso: {titulo_curso}")
        detalhes_janela.geometry("400x300")

        # Exibir o título centralizado
        tk.Label(detalhes_janela, text="Detalhes do Curso", font=("Arial", 20, "bold")).pack(pady=10, anchor="center")

        # Criar um frame para os detalhes e alinhar com margem na esquerda
        frame_detalhes = tk.Frame(detalhes_janela)
        frame_detalhes.pack(fill="both", expand=True, padx=20, pady=5)

        # Exibir os detalhes do curso com margem à esquerda
        tk.Label(frame_detalhes, text=f"Título: {titulo_curso}", font=("Arial", 14), anchor="w").pack(pady=5, fill="x")
        tk.Label(frame_detalhes, text=f"Nota: {nota}", font=("Arial", 14), anchor="w").pack(pady=5, fill="x")
        tk.Label(frame_detalhes, text=f"Presença: {presenca}%", font=("Arial", 14), anchor="w").pack(pady=5, fill="x")
        tk.Label(frame_detalhes, text=f"Certificado: {certificado}", font=("Arial", 14), anchor="w").pack(pady=5,
                                                                                                          fill="x")

        # Botão para fechar a janela, alinhado ao centro
        tk.Button(detalhes_janela, text="Fechar", font=("Arial", 14), command=detalhes_janela.destroy).pack(pady=20)

    except Exception as e:
        messagebox.showerror("Erro ao Abrir Curso", f"Ocorreu um erro ao tentar abrir os detalhes do curso.\n\n{e}")

    def consultar_cursos_inscritos(tree, id_perfis):
        """
        Consulta os cursos em que o usuário está inscrito e exibe na Treeview.
        """
        try:
            # Conectar ao banco de dados
            conn = conexaoBanco()
            cursor = conn.cursor()

            # Consulta SQL para buscar os cursos inscritos pelo usuário
            cursor.execute("""
                SELECT c.id_curso, c.nome, c.carga_horaria, c.vagas
                FROM inscricoes i
                JOIN cursos c ON i.id_curso = c.id_curso
                WHERE i.id_perfis = %s
            """, (id_perfis,))

            # Limpar o Treeview antes de inserir novos dados
            for item in tree.get_children():
                tree.delete(item)

            # Processar os resultados e adicionar ao Treeview
            for row in cursor.fetchall():
                id, titulo, carga_horaria, vagas = row
                disponibilidade = "Disponível" if vagas > 0 else "Indisponível"
                tree.insert("", "end", values=(id, titulo, carga_horaria, disponibilidade))

            # Fechar conexão
            cursor.close()
            conn.close()

        except Exception as e:
            messagebox.showerror("Erro ao Consultar", f"Não foi possível consultar os cursos inscritos.\n\n{e}")

def emitir_certificado(tree, perfil):
    """
    Gera um certificado do curso selecionado no TreeView e o salva na pasta 'certificados'.
    """
    try:
        # Obter o item selecionado no TreeView
        selected_item = tree.focus()
        if not selected_item:
            messagebox.showwarning("Seleção Inválida", "Por favor, selecione um curso.")
            return

        # Obter os valores da linha selecionada
        valores = tree.item(selected_item, "values")
        if not valores:
            messagebox.showerror("Erro", "Erro ao obter os dados do curso.")
            return

        # Extrair os detalhes do curso
        titulo_curso, nota, presenca, certificado = valores

        # Verificar se o certificado está disponível
        if certificado != "Disponível":
            messagebox.showerror("Certificado Indisponível", "Você não atingiu os requisitos para obter o certificado.")
            return

        # Diretório de saída para os certificados
        diretorio_certificados = os.path.join(os.getcwd(), "certificados")
        if not os.path.exists(diretorio_certificados):
            os.makedirs(diretorio_certificados)  # Criar a pasta se ela não existir

        # Criar o documento do Word
        doc = Document()
        doc.add_heading('Certificado de Conclusão', level=1)
        doc.add_paragraph(
            f"Certificamos que {perfil} concluiu com êxito o curso '{titulo_curso}', "
            f"com a nota final de {nota} e presença de {presenca}%."
        )
        doc.add_paragraph(
            f"Data: {datetime.now().strftime('%d/%m/%Y')}\n\n"
            "Este certificado foi gerado automaticamente e é válido somente para fins acadêmicos."
        )

        # Salvar o arquivo na pasta 'certificados'
        arquivo_certificado = os.path.join(
            diretorio_certificados,
            f"Certificado_{titulo_curso.replace(' ', '_')}.docx"
        )
        doc.save(arquivo_certificado)

        messagebox.showinfo("Certificado Emitido", f"Certificado emitido com sucesso: {arquivo_certificado}")

    except Exception as e:
        messagebox.showerror("Erro ao Emitir Certificado", f"Ocorreu um erro ao tentar emitir o certificado.\n\n{e}")

def consultar_cursos_inscritos(tree, login_usuario):
    """
    Consulta os cursos em que o usuário está inscrito, utilizando o login, e exibe na Treeview.
    """
    try:
        # Conectar ao banco de dados
        conn = conexaoBanco()
        cursor = conn.cursor()

        # Consulta SQL para buscar os cursos inscritos pelo usuário e calcular a presença
        query = """
            SELECT c.nome AS titulo, 
                   i.nota_aluno, 
                   ROUND((i.presenca_aluno / c.carga_horaria) * 100, 0) AS presenca,
                   CASE 
                       WHEN i.nota_aluno >= 7 AND (i.presenca_aluno / c.carga_horaria) * 100 >= 80 THEN 'Disponível'
                       ELSE 'Indisponível'
                   END AS certificado
            FROM inscricoes i
            JOIN cursos c ON i.id_curso = c.id_curso
            JOIN perfis p ON i.id_perfis = p.id_perfis
            WHERE p.login = %s
        """
        cursor.execute(query, (login_usuario,))

        # Limpar o Treeview antes de inserir novos dados
        for item in tree.get_children():
            tree.delete(item)

        # Processar os resultados e adicionar ao Treeview
        resultados = cursor.fetchall()
        print(f"Resultados da consulta: {resultados}")  # Debug para verificar os dados retornados

        if not resultados:
            messagebox.showinfo("Sem Dados", "Nenhum curso encontrado para este usuário.")
            return

        for row in resultados:
            titulo, nota, presenca, certificado = row
            tree.insert("", "end", values=(titulo, nota, presenca, certificado))

        # Fechar conexão
        cursor.close()
        conn.close()

    except Exception as e:
        messagebox.showerror("Erro ao Consultar", f"Não foi possível consultar os cursos inscritos.\n\n{e}")

def iniciar(perfil):

    root = tk.Tk()
    root.title("Home")
    root.geometry("700x400")
    root.resizable(True, True)

    # Configurando o frame superior para usar grid
    top_frame = tk.Frame(root, height=50)
    top_frame.pack(side="top", fill="x")

    # Configurar o grid dentro do `top_frame`
    top_frame.grid_columnconfigure(0, weight=1)  # Espaço antes do logo
    top_frame.grid_columnconfigure(2, weight=1)  # Espaço central
    top_frame.grid_columnconfigure(4, weight=1)  # Espaço após o botão

    # Imagem CEUB (na parte esquerda do grid)
    diretorio_atual = os.path.dirname(__file__)
    caminho_imagem = os.path.join(diretorio_atual, "..", "..", "imagens", "uniceub.png")
    logo_img = PhotoImage(file=caminho_imagem)
    logo_label = tk.Label(top_frame, image=logo_img)
    logo_label.grid(row=0, column=0, padx=10, sticky="w")  # Alinhado à esquerda

    # Campo de busca (no centro do grid)
    titulo = tk.Label(top_frame, text="Gestão de Cursos Monitoria", font='Helvetica 16 bold')
    titulo.grid(row=0, column=2, padx=20, sticky="ew")  # Expansível horizontalmente

    # Botão Nome do usuário (na direita do grid)
    botao_usuario = tk.Button(top_frame, text="Meu Perfil", font=("Arial", 12), command=lambda: abrir_dados_usuario(perfil))
    botao_usuario.grid(row=0, column=4, padx=20, sticky="e")  # Alinhado à direita

    # Barra lateral
    sidebar = tk.Frame(root, width=200)
    sidebar.pack(side="left", fill="y")

    # Botões do menu lateral
    btn_home = tk.Button(sidebar, text="Home", font=("Arial", 10, "bold"), relief="flat", bg="lightblue", fg="White")
    btn_home.pack(pady=10, padx=10, fill="x")

    separator = tk.Frame(sidebar, height=2, bg="white")
    separator.pack(fill="x", pady=10, padx=10)  # Linha separadora

    btn_meus_cursos = tk.Button(sidebar, text="Meus Cursos", font=("Arial", 10, "bold"), relief="flat", command=lambda: meusCursos(perfil))
    btn_meus_cursos.pack(pady=10, padx=10, fill="x")

    btn_descobrir = tk.Button(sidebar, text="Descobrir", font=("Arial", 10, "bold"), relief="flat", command=lambda: descobrir(perfil))
    btn_descobrir.pack(pady=10, padx=10, fill="x")

    separator = tk.Frame(sidebar, height=2, bg="white")
    separator.pack(fill="x", pady=10, padx=10)  # Linha separadora

    # Área de conteúdo principal
    content_area = tk.Frame(root)
    content_area.pack(side="right", expand=True, fill="both")

    # Mensagem de boas-vindas no topo esquerdo da área de conteúdo
    # Ajustando o label para exibir o nome do usuário dinamicamente
    label_boas_vindas = tk.Label(content_area, text=f"Seja bem-vindo, {perfil}!",font="Arial 14",anchor="w",padx=10)
    label_boas_vindas.grid(row=0, column=0, padx=10, pady=10)

    # Botão de consulta no topo direito da área de conteúdo
    bt_consultar = tk.Button(content_area, text="Consultar", command=lambda: consultar_cursos_inscritos(tree, perfil))
    bt_consultar.grid(row=0, column=2, padx=10, pady=10)

    # Adicionando o Treeview na área de conteúdo logo abaixo
    tree = ttk.Treeview(content_area, columns=("Título", "Nota", "Presença", "Certificado"), show="headings")
    tree.heading("Título", text="Título do Curso")
    tree.heading("Nota", text="Nota")
    tree.heading("Presença", text="Presença (%)")
    tree.heading("Certificado", text="Certificado")
    tree.column("Título", width=250)
    tree.column("Nota", width=80, anchor="center")
    tree.column("Presença", width=100, anchor="center")
    tree.column("Certificado", width=100, anchor="center")
    tree.grid(row=1, column=0, columnspan=3, padx=10, pady=10)

    bt_abrir = tk.Button(content_area, text="Abrir", command=lambda: abrir_curso(tree))
    bt_abrir.grid(row=2, column=2, padx=10, pady=10)

    bt_incluir = tk.Button(content_area, text="Emitir certificado", command=lambda: emitir_certificado(tree, perfil))
    bt_incluir.grid(row=2, column=0, padx=10, pady=10)

    root.mainloop()